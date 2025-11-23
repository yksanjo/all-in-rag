"""
Document loader for various file formats.
Supports PDF, DOCX, TXT, and image files.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from PIL import Image


class BaseDocumentLoader(ABC):
    """Base class for document loaders."""
    
    @abstractmethod
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with 'content', 'metadata', and 'type' keys
        """
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """Check if this loader supports the given file."""
        pass


class PDFLoader(BaseDocumentLoader):
    """Loader for PDF files."""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load PDF document."""
        if not HAS_PDF:
            raise ImportError("PyPDF2 is required for PDF support")
        
        content = []
        metadata = {
            "source": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "pdf"
        }
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata["num_pages"] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    content.append({
                        "page": page_num + 1,
                        "text": text
                    })
        
        return {
            "content": content,
            "metadata": metadata,
            "type": "text"
        }
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a PDF."""
        return file_path.lower().endswith('.pdf') and HAS_PDF


class DOCXLoader(BaseDocumentLoader):
    """Loader for DOCX files."""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load DOCX document."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX support")
        
        doc = Document(file_path)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                table_text.append(row_text)
            if table_text:
                content.append("\n".join(table_text))
        
        metadata = {
            "source": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "docx"
        }
        
        return {
            "content": "\n\n".join(content),
            "metadata": metadata,
            "type": "text"
        }
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return file_path.lower().endswith('.docx') and HAS_DOCX


class TXTLoader(BaseDocumentLoader):
    """Loader for plain text files."""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load text document."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Could not decode file: {file_path}")
        
        metadata = {
            "source": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "txt"
        }
        
        return {
            "content": content,
            "metadata": metadata,
            "type": "text"
        }
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a text file."""
        return file_path.lower().endswith(('.txt', '.md', '.rst'))


class ImageLoader(BaseDocumentLoader):
    """Loader for image files."""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load image document."""
        try:
            image = Image.open(file_path)
            metadata = {
                "source": file_path,
                "file_name": os.path.basename(file_path),
                "file_type": "image",
                "width": image.width,
                "height": image.height,
                "format": image.format
            }
            
            return {
                "content": file_path,  # Store path, image will be loaded when needed
                "metadata": metadata,
                "type": "image"
            }
        except Exception as e:
            raise ValueError(f"Could not load image: {file_path}, error: {e}")
    
    def supports(self, file_path: str) -> bool:
        """Check if file is an image."""
        image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
        return Path(file_path).suffix.lower() in image_extensions


class DocumentLoader:
    """Main document loader that routes to appropriate loader."""
    
    def __init__(self):
        """Initialize document loader with all available loaders."""
        self.loaders: List[BaseDocumentLoader] = [
            PDFLoader(),
            DOCXLoader(),
            TXTLoader(),
            ImageLoader()
        ]
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with document content and metadata
            
        Raises:
            ValueError: If file format is not supported
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        for loader in self.loaders:
            if loader.supports(file_path):
                return loader.load(file_path)
        
        raise ValueError(f"Unsupported file format: {file_path}")
    
    def load_directory(self, directory: str, recursive: bool = True) -> List[Dict[str, Any]]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Path to directory
            recursive: Whether to search recursively
            
        Returns:
            List of loaded documents
        """
        documents = []
        path = Path(directory)
        
        if recursive:
            files = path.rglob('*')
        else:
            files = path.glob('*')
        
        for file_path in files:
            if file_path.is_file():
                try:
                    doc = self.load(str(file_path))
                    documents.append(doc)
                except (ValueError, ImportError) as e:
                    # Skip unsupported files
                    continue
        
        return documents

